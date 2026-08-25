# -*- coding: utf-8 -*-
"""공식 결과보고서 docx 칸 채우기. 일회성 제출 보조 스크립트."""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path(
    r"c:\Users\wjsto\Downloads\46414fba-c473-4dae-b595-7214d635b494 (1)"
    r"\2026 오픈소스 개발자대회 결과보고서_접수번호(팀명)"
    r"\2026 오픈소스 개발자대회 결과보고서_915(지엔).docx"
)
BACKUP = SRC.with_name(SRC.stem + "_빈양식백업.docx")
REPO_OUT = Path(
    r"C:\Users\wjsto\pjt\capnet\docs\ops"
    r"\2026 오픈소스 개발자대회 결과보고서_915(지엔).docx"
)


def set_text(cell, text: str) -> None:
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)
    lines = text.split("\n") if text else [""]
    for line in lines:
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r.append(t)
        p.append(r)
        tc.append(p)


def unique_cells(row):
    seen: set[int] = set()
    out = []
    for cell in row.cells:
        tid = id(cell._tc)
        if tid in seen:
            continue
        seen.add(tid)
        out.append(cell)
    return out


def add_table_row(table):
    tbl = table._tbl
    last = table.rows[-1]._tr
    new_tr = deepcopy(last)
    for t in new_tr.iter(qn("w:t")):
        t.text = ""
    tbl.append(new_tr)
    return table.rows[-1]


def main() -> None:
    if not BACKUP.exists():
        BACKUP.write_bytes(SRC.read_bytes())

    d = Document(str(SRC))

    # 안내 배너 비우기
    set_text(d.tables[0].rows[0].cells[0], "")

    # 팀 정보 — 인원·부문은 접수와 다르면 Word에서 수정
    t2 = d.tables[2]
    set_text(t2.rows[1].cells[1], "지엔")
    set_text(t2.rows[1].cells[3], "3명")
    set_text(t2.rows[2].cells[1], "일반")
    set_text(t2.rows[2].cells[3], "자유과제")

    t3 = d.tables[3]
    set_text(t3.rows[1].cells[1], "CapNet")
    set_text(t3.rows[2].cells[1], "https://github.com/gncorpseo-commits/capnet")
    set_text(t3.rows[3].cells[1], "https://youtu.be/RjFiGpmLTbk")
    set_text(
        t3.rows[4].cells[1],
        "CapNet은 AI 모델을 공유하는 서비스가 아니라, 여러 실행 자원을 하나의 실행 계층으로 묶고 "
        "사용자는 서버 주소가 아니라 필요한 AI 능력만 요청하도록 만드는 오픈소스 실행 계층이다. "
        "Core가 승인된 조건 안에서만 배정하고, 어떤 Agent·Node에서 실행됐는지 증적을 남긴다.",
    )
    set_text(
        t3.rows[6].cells[1],
        "대부분의 AI 호출은 GPU 서버 주소 → 추론 API → 결과 순서다. 결과만 맞아도, 규정·감사·데이터 통제가 "
        "필요한 조직에서는 무엇이 답했는지·자격이 있었는지·데이터가 어디까지 갔는지를 사후에 확인하기 어렵다.\n"
        'CapNet은 내 서비스 → "image.classify" → Core → 조건에 맞는 실행기 → 결과 구조로 바꾼다. '
        "GPU를 찾는 방법이 아니라 이 일을 할 실행 자원을 누가 결정하느냐가 핵심이다. "
        "Core는 신뢰 도메인·연산 등급·게이트 통과 Agent 조건 안에서만 배정하고, "
        "잘못된 조합은 애플리케이션이 아니라 PostgreSQL이 거절한다.",
    )
    set_text(
        t3.rows[7].cells[1],
        "- HW: 팀 노트북 CPU · Docker Desktop / WSL2\n"
        "- SW: Python 3.11 · PostgreSQL 16 · FastAPI · PyTorch 2.13.0+cpu · torchvision 0.28.0+cpu · safetensors\n"
        "- 실행: Docker Compose (postgres 1 + Core 1 + Node 3)\n"
        "- 데이터: EuroSAT RGB (Zenodo 7711810, MIT) — 원본 미동봉, 다운로드 스크립트 제공\n"
        "- 모델: 사전학습 가중치 미사용. EuroSAT scratch 학습만\n"
        "- 외부 상용 AI API 호출 없음 — 추론 전량 로컬 컨테이너",
    )
    set_text(
        t3.rows[8].cells[1],
        "- Node: 실행 가능한 컴퓨터. Core가 승인한 작업만 수행 (자기 등급 주장 불가)\n"
        "- Agent: 특정 AI 능력의 구현\n"
        "- Capability: 사용자가 요구하는 작업의 계약\n"
        "- Core(:8000): Capability × Agent × Node 조합으로 배정·증적. 사용자 접점\n"
        "- postgres: 라우팅·게이트 불변식(복합 FK·호환 행렬·증서)\n"
        "- 연결: Node → Core outbound(HTTP polling). WebSocket은 향후 확장\n"
        "- 흐름: 사용자→Core(능력만 요청) → Core 배정 → Node가 배정 가져와 실행 → 결과·증적 조회\n"
        "- 무단 호출 거부(HTTP 403). 사용자는 기기 주소를 모른다",
    )
    set_text(
        t3.rows[9].cells[1],
        "【프로젝트 상세 내용】\n"
        "실행 자원 네트워크. 사용자는 Node를 지정하지 않고 능력만 요청한다. 결과뿐 아니라 경로(증적)가 제품이다.\n"
        "1) 능력=계약 — 입출력 스키마·전처리·실행 조건. 6종 능력이 같은 사슬 통과"
        "(image.classify, text.classify, image.embed, text.embed, timeseries.forecast, table.extract).\n"
        "2) DB가 라우팅 강제 — privacy_rank 등 복합 FK. team 작업→public Node 배정 불가.\n"
        "3) 살아 있는 Node만 — liveness·부하 기반 배정.\n"
        "4) 실행 증적 — node·agent·가중치 해시·시각을 DB에 기록·조회.\n"
        "5) 입력 통제 — Core 중개·해시·MIME 검증. 자유 업로드 경로 없음.\n"
        "6) 위반 6종 DB 거절 — 자격 없는 배정, 도메인·등급 위반, lease 중 강등, READY 중 가중치 교체, PASSED 무효화 등.\n"
        "7) 품질 프로파일(선택) — image.classify@1 scratch 실측 acc=0.8500 · macro-F1=0.8344 PASSED. "
        "상수·난수·스키마 위반 Agent는 전부 FAILED.\n"
        "\n"
        "【구동 및 시연】\n"
        "git clone https://github.com/gncorpseo-commits/capnet.git\n"
        "cd capnet && docker compose up --build -d\n"
        "bash scripts/demo.sh && bash scripts/sanity.sh && bash scripts/demo_violations.sh\n"
        "(Windows: 동명 .ps1) 기대 출력: PASSED acc=0.8500 · sanity 3× FAILED · violations 6× REJECTED.\n"
        "demo에 기기 주소 없음 — 사용자는 Core만 호출한다.",
    )
    set_text(
        t3.rows[10].cells[1],
        "【향후 확장성 및 기대효과】\n"
        "단기(출품 범위): 조직 내부의 신뢰된 Node 플릿에서 실행 자원 네트워크가 성립함을 증명.\n"
        "확장: 유휴 PC를 풀로 묶되 “team 영역만·M급 이상·gate 통과 Agent만” 조건을 Core가 적용. "
        "PC를 서버로 직접 관리하는 방식에서 주소가 아닌 능력으로 실행 자원을 쓰는 방식으로 전환.\n"
        "값은 “AI 스토어”가 아니라 실행 계층과 증적에 있다. "
        "데이터를 외부로 보낼 수 없는 조직이 자체 장비에서 AI 작업을 처리하되 경로를 남기는 구조에 적용된다.",
    )
    set_text(
        t3.rows[11].cells[1],
        "【프로젝트의 혁신성 및 차별성】\n"
        "CapNet은 AI 서버를 하나 더 만드는 프로젝트가 아니다. "
        "라우팅 불변식을 애플리케이션이 아니라 PostgreSQL이 강제한다. "
        "위반 14종을 실측하고 그중 6종을 재현 스크립트로 고정했다. "
        "사용자는 기기에 직접 접속하지 않으며, 입력 바이트도 Core가 중개한다.\n"
        "\n"
        "【한계점 및 향후 발전 로드맵】\n"
        "기기가 데이터를 남기지 않는다는 보장은 하지 않는다(TEE 없이 원리적으로 불가). "
        "보장하는 것은 승인하지 않은 신뢰 도메인으로 라우팅되지 않는다는 것과 증적이 남는다는 것이다. "
        "현재는 공개 실행 네트워크가 아니라 조직 내부 신뢰 플릿 검증 단계다. "
        "골든셋 정적·공개 한계와 A/B 등가 한계는 문서화했다. "
        "다음: 홀드아웃·회전 프로브·조직 단위 플릿.\n"
        "\n"
        "【소감 및 후기】\n"
        "제약을 끄지 않고 막힐 때마다 DB를 믿도록 설계를 고쳤다. "
        "골든셋 누출·A/B 등가·직접 Node 호출 문제를 스스로 발견해 문서에 남겼다. "
        "파이프라인이 동작한다는 것과 측정이 유효하다는 것은 다른 명제다.",
    )

    sbom = [
        ("1", "psycopg", "3.2.9", "LGPL-3.0", "https://github.com/psycopg/psycopg", "PostgreSQL 드라이버"),
        ("2", "psycopg-pool", "3.3.1", "LGPL-3.0", "https://github.com/psycopg/psycopg", "커넥션 풀"),
        ("3", "FastAPI", "0.116.1", "MIT", "https://github.com/tiangolo/fastapi", "Core·Node HTTP API"),
        ("4", "uvicorn", "0.35.0", "BSD-3-Clause", "https://github.com/encode/uvicorn", "ASGI 서버"),
        ("5", "pydantic-settings", "2.10.1", "MIT", "https://github.com/pydantic/pydantic-settings", "환경변수 설정"),
        ("6", "safetensors", "0.8.0", "Apache-2.0", "https://github.com/huggingface/safetensors", "가중치 로드(pickle 거부)"),
        ("7", "numpy", "2.4.6", "BSD-3-Clause", "https://github.com/numpy/numpy", "텐서·배열 백엔드"),
        ("8", "Pillow", "12.3.0", "HPND-derived", "https://github.com/python-pillow/Pillow", "골든셋 JPEG 로드"),
        ("9", "torch", "2.13.0+cpu", "BSD-3-Clause", "https://github.com/pytorch/pytorch", "scratch 학습·추론"),
        ("10", "torchvision", "0.28.0+cpu", "BSD-3-Clause", "https://github.com/pytorch/vision", "32×32 변환"),
    ]
    t5 = d.tables[5]
    while len(t5.rows) < 11:
        add_table_row(t5)
    for i, row_data in enumerate(sbom):
        row = t5.rows[i + 1]
        for ci, val in enumerate(row_data):
            set_text(row.cells[ci], val)
    if len(t5.rows) > 11:
        for ri in range(11, len(t5.rows)):
            for cell in unique_cells(t5.rows[ri]):
                set_text(cell, "")

    set_text(d.tables[7].rows[0].cells[0], "")

    t8 = d.tables[8]
    set_text(
        unique_cells(t8.rows[1])[0],
        "□ 유형 1: 외부 모델 그대로 활용 (추가 학습 없이 기존 공개 모델을 프로젝트에 연동·구동한 경우)\n"
        "□ 유형 2: 외부 모델 파인튜닝 (기존 공개 모델을 가져와 준비한 데이터셋으로 추가 미세조정한 경우)\n"
        "▣ 유형 3: 자체 개발 모델 (기반 모델 없이 참가팀이 처음부터 가중치를 직접 전체 학습시킨 경우)\n"
        "판정: 기반 모델 없이 EuroSAT RGB만으로 가중치를 처음부터 학습 → 유형 3. "
        "외부 상용 AI API 의존 없음(로컬 torch). 코딩 보조는 아래 4번 항목에 기재.",
    )
    uc3 = unique_cells(t8.rows[3])
    set_text(uc3[1], "해당 없음 (기반 모델 미사용. TinyEuroSAT을 EuroSAT RGB로 scratch 학습)")
    set_text(uc3[3], "해당 없음")
    set_text(
        unique_cells(t8.rows[5])[1],
        "EuroSAT RGB 배포판 · Zenodo record 7711810 · DOI 10.5281/zenodo.7711810 · 라이선스 MIT · "
        "EuroSAT_RGB.zip 27,000장(10클래스) · 원본 64×64 JPEG. "
        "원본 zip은 저장소 미동봉(scripts/download_eurosat + archive_sha256 핀).",
    )
    set_text(
        unique_cells(t8.rows[6])[1],
        "개인정보 없음(원격탐사 타일). 클래스 균등 스트라이드로 데모 골든 N=40 추출. "
        "게이트·제품 계약은 32×32 resize. 학습은 전체 27,000장 · seed 20260806 · "
        "Agent A 80 epoch · Agent B 40 epoch(CPU, scratch, pretrained=false).",
    )
    set_text(
        unique_cells(t8.rows[7])[1],
        "https://github.com/gncorpseo-commits/capnet/blob/main/apps/node/weights/eurosat_scratch.safetensors\n"
        "https://github.com/gncorpseo-commits/capnet/blob/main/apps/node/weights/eurosat_scratch_b.safetensors\n"
        "직접 다운로드(승인 없음):\n"
        "https://raw.githubusercontent.com/gncorpseo-commits/capnet/main/apps/node/weights/eurosat_scratch.safetensors\n"
        "https://raw.githubusercontent.com/gncorpseo-commits/capnet/main/apps/node/weights/eurosat_scratch_b.safetensors",
    )
    set_text(
        unique_cells(t8.rows[8])[1],
        "eurosat_scratch.safetensors / 전체 가중치 배포 (TinyEuroSAT CNN, 94,538 파라미터) / 370KB · "
        "sha256 74ca92224ff93f6cfab56265466d2c8ed11e0add4581c45a98169e92fb797b43\n"
        "eurosat_scratch_b.safetensors / 전체 가중치 배포 (TinyEuroSATB CNN, 24,685 파라미터) / 98KB · "
        "sha256 3fbdde549459ba1e895ab8221ba1455b7840e1641b36b202c96420d410343b24\n"
        "LoRA가 아닌 전체 가중치(유형 3). 메타는 각 *.meta.json.",
    )
    uc10 = unique_cells(t8.rows[10])
    set_text(uc10[1], "Apache License 2.0")
    set_text(
        uc10[3],
        "https://github.com/gncorpseo-commits/capnet\n"
        "(apps/train/train_scratch.py, apps/node/app/infer.py, apps/node/app/score_gate.py)",
    )
    set_text(
        unique_cells(t8.rows[11])[1],
        "코드 작성·디버깅·문서 정리 보조로 Cursor 및 Claude Code를 활용함. "
        "전체 코드(약 16,903줄, py·sh·ps1·sql·yaml 기준) 중 AI 보조 작성분은 약 30%로 추정함. "
        "스키마 제약·게이트 사슬·채점 규칙·판정 기준은 팀이 설계했고, "
        "AI 생성 코드는 팀이 동작 원리를 확인하고 실행 검증한 뒤 반영함.",
    )

    d.save(str(SRC))
    REPO_OUT.write_bytes(SRC.read_bytes())
    print("saved", SRC)
    print("backup", BACKUP)
    print("repo_copy", REPO_OUT)
    print("bytes", SRC.stat().st_size)


if __name__ == "__main__":
    main()
